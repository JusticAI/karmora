import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'PingFang SC'
style.font.size = Pt(11)

def title(t):
    p = doc.add_heading(t, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0x45,0x00); r.font.size = Pt(24)

def sub(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x78,0x7C,0x7E)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0x45,0x00); r.font.size = Pt(18)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.size = Pt(14)

def h3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs: r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

def b(t):
    doc.add_paragraph(t).paragraph_format.space_after = Pt(6)

def li(t):
    doc.add_paragraph(t, style='List Bullet')

def code(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.font.name = 'Menlo'
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x33,0x33,0x33)

def div():
    p = doc.add_paragraph()
    r = p.add_run('─'*60); r.font.color.rgb = RGBColor(0xDD,0xDD,0xDD); r.font.size = Pt(8)

# ═══════════════════════════════════════════
title('Karmora 数据获取方法汇总')
sub('项目开发过程中发现的所有 Reddit 数据获取方案')
sub('用于后期社区数据扩展')
sub('生成日期：2026-04-21')
doc.add_page_break()

# ── 总览 ──
h1('数据获取方案总览')
b('在项目开发过程中，共发现 6 种 Reddit 数据获取方法。以下按可行性和实用性排序：')
doc.add_paragraph()

# Table
table = doc.add_table(rows=7, cols=5)
table.style = 'Light Grid Accent 1'
headers = ['方法', '可行性', '需要条件', '数据类型', '限制']
for i, h in enumerate(headers):
    table.rows[0].cells[i].text = h
rows_data = [
    ['① 浏览器同源请求', '✅ 最佳', '用户登录 Reddit', '帖子/评论/用户/社区', '需登录'],
    ['② PullPush API', '✅ 可用', '无', '历史帖子/评论', '慢(3-4s)，超时'],
    ['③ Reddit about/rules', '✅ 可用', '用户登录', '社区规则/描述', '需浏览器内'],
    ['④ Wiki 页面抓取', '✅ 可用', '用户登录', '社区详细规则', 'CSS选择器不稳定'],
    ['⑤ Reddit 官方 API', '❌ 阻断', 'Client ID+Secret', '全部', '2023年起封服务器端'],
    ['⑥ 浏览器直接抓取', '⚠️ 不稳定', '无', '页面可见内容', 'Cloudflare/CAPTCHA'],
]
for i, row in enumerate(rows_data):
    for j, cell in enumerate(row):
        table.rows[i+1].cells[j].text = cell

doc.add_page_break()

# ── 方法1 ──
h1('方法①：浏览器同源请求（最佳方案）')

h2('原理')
b('当用户在浏览器中登录 Reddit 后，从 Reddit 页面的 Console 中发起的 fetch() 请求会自动携带 cookies（同源策略）。Reddit 封锁的是"外部服务器直接调 API"，但"从浏览器内部发请求"不受影响。')

h2('使用场景')
li('Chrome 扩展的 content script（在 reddit.com 域名下运行）')
li('用户在 Reddit 页面手动打开 DevTools Console 执行代码')

h2('可用的 API 端点')

h3('获取当前用户信息')
code('const resp = await fetch("/api/me.json", {credentials: "same-origin"});')
code('const data = await resp.json();')
code('// 返回: name, link_karma, comment_karma, created_utc, ...')

h3('获取社区帖子列表')
code('const resp = await fetch("/r/{subreddit}/new.json?limit=50", {credentials: "same-origin"});')
code('const data = await resp.json();')
code('// 返回: 50条最新帖子，包含 title, selftext, num_comments, score, created_utc, ...')

h3('获取社区帖子（按热度）')
code('const resp = await fetch("/r/{subreddit}/hot.json?limit=50", {credentials: "same-origin"});')

h3('获取社区帖子（按评分）')
code('const resp = await fetch("/r/{subreddit}/top.json?t=month&limit=50", {credentials: "same-origin"});')

h3('获取社区信息')
code('const resp = await fetch("/r/{subreddit}/about.json", {credentials: "same-origin"});')
code('// 返回: description, subscribers, created_utc, public_description, ...')

h3('获取社区规则')
code('const resp = await fetch("/r/{subreddit}/rules.json", {credentials: "same-origin"});')
code('// 返回: 规则列表，每条包含 short_name, description, violation_reason, ...')

h3('获取社区 Wiki 页面列表')
code('const resp = await fetch("/r/{subreddit}/wiki/pages.json", {credentials: "same-origin"});')
code('// 返回: wiki 页面名称列表')

h3('获取社区 Wiki 页面内容')
code('const resp = await fetch("/r/{subreddit}/wiki/{page}.json", {credentials: "same-origin"});')
code('// 返回: content_md (Markdown格式的wiki内容)')

h3('搜索社区')
code('const resp = await fetch("/subreddits/search.json?q={keyword}&limit=50", {credentials: "same-origin"});')
code('// 返回: 匹配的社区列表')

h3('搜索社区内帖子')
code('const resp = await fetch("/r/{subreddit}/search.json?q={keyword}&restrict_sr=on&limit=50", {credentials: "same-origin"});')

h3('获取用户帖子历史')
code('const resp = await fetch("/user/{username}/submitted.json?limit=50", {credentials: "same-origin"});')

h3('获取用户评论历史')
code('const resp = await fetch("/user/{username}/comments.json?limit=50", {credentials: "same-origin"});')

h2('已知问题')
li('只在用户登录 Reddit 的浏览器内有效')
li('rate limit: Reddit 对同源请求也有频率限制，太快会被临时封锁')
li('返回最多 100 条/请求，需要分页获取更多')

h2('在 Karmora 中的应用')
li('Chrome 扩展 content.js 中的 RedditAPI 模块')
li('用户在 Console 手动运行数据采集脚本')
li('采集 40 个社区的 engagement 数据（avg_comments, zero_rate, posts_per_day）')

div()

# ── 方法2 ──
h1('方法②：PullPush API（批量历史数据）')

h2('原理')
b('PullPush（原 Pushshift）是一个 Reddit 数据归档服务，提供免费 API 查询历史帖子和评论。不需要登录，不需要 API 凭证。')

h2('API 端点')

h3('搜索社区帖子')
code('GET https://api.pullpush.io/reddit/search/submission/')
code('?subreddit={subreddit}&sort=score&order=desc&size=50')
code('// 返回: 该社区的帖子列表（按评分排序）')

h3('搜索社区帖子（按时间）')
code('?subreddit={subreddit}&sort=created_utc&order=desc&size=50')

h3('搜索特定关键词')
code('?subreddit={subreddit}&q={keyword}&sort=score&order=desc&size=50')

h3('搜索评论')
code('GET https://api.pullpush.io/reddit/search/comment/')
code('?subreddit={subreddit}&sort=score&order=desc&size=50')

h2('Python 调用示例')
code('import urllib.request, json, time')
code('')
code('subs = ["TooAfraidToAsk", "CasualConversation", ...]')
code('for sub in subs:')
code('    url = f"https://api.pullpush.io/reddit/search/submission/?subreddit={sub}&sort=score&order=desc&size=50"')
code('    req = urllib.request.Request(url, headers={"User-Agent": "KarmoraResearch/1.0"})')
code('    with urllib.request.urlopen(req, timeout=15) as resp:')
code('        data = json.loads(resp.read())')
code('        posts = data.get("data", [])')
code('        # 计算指标...')
code('    time.sleep(1)  # 必须加延迟，否则被封')

h2('已知问题')
li('慢：每个请求 3-4 秒')
li('搜索查询经常超时：不要用复杂查询，只用基础 subreddit 端点')
li('数据可能不是最新的：归档数据有延迟')
li('需要 User-Agent header，否则返回 403')
li('批量采集需要加 0.5-1 秒延迟')

h2('在 Karmora 中的应用')
li('Python 脚本批量采集 131 个社区的数据')
li('计算 engagement 指标')

div()

# ── 方法3 ──
h1('方法③：Reddit about.json + rules.json（社区元数据）')

h2('原理')
b('每个 subreddit 都有 about.json 和 rules.json 两个公开端点，返回社区的基本信息和规则列表。')

h2('about.json 返回数据')
code('{')
code('  "data": {')
code('    "title": "r/ecommerce",')
code('    "public_description": "A community for ecommerce professionals...",')
code('    "subscribers": 123456,')
code('    "accounts_active": 1234,')
code('    "created_utc": 1234567890,')
code('    "over18": false,')
code('    "lang": "en",')
code('    "subreddit_type": "public",')
code('    "wiki_enabled": true,')
code('    ...')
code('  }')
code('}')

h2('rules.json 返回数据')
code('{')
code('  "rules": [')
code('    {')
code('      "short_name": "No self-promotion",')
code('      "description": "Do not promote your own business...",')
code('      "violation_reason": "Spam",')
code('      "created_utc": 1234567890,')
code('      "priority": 0')
code('    },')
code('    ...')
code('  ],')
code('  "sitewide_rules": [...]')
code('}')

h2('使用方法')
code('// 从浏览器 Console 或 Chrome 扩展')
code('const about = await fetch("/r/ecommerce/about.json", {credentials: "same-origin"});')
code('const rules = await fetch("/r/ecommerce/rules.json", {credentials: "same-origin"});')

h2('价值')
li('社区规则是 AI 助手"规则合规"功能的核心数据')
li('社区描述用于智能匹配')
li('subscribers 用于判断社区规模')

div()

# ── 方法4 ──
h1('方法④：Wiki 页面抓取（详细规则）')

h2('原理')
b('很多社区在 Wiki 页面中维护详细的发帖指南、格式要求、常见问题等。这些信息比 rules.json 更详细。')

h2('获取 Wiki 页面列表')
code('const resp = await fetch("/r/{sub}/wiki/pages.json", {credentials: "same-origin"});')
code('// 返回: ["index", "faq", "rules", "guidelines", ...]')

h2('获取 Wiki 页面内容')
code('const resp = await fetch("/r/{sub}/wiki/{page}.json", {credentials: "same-origin"});')
code('const data = await resp.json();')
code('const markdown = data.data.content_md;  // Markdown 格式')

h2('已知问题')
li('CSS 选择器 .wiki-page-content 不匹配所有页面')
li('需要 fallback 链：.wiki-page-content → .md.wiki → .md')
li('不是所有社区都启用 Wiki')

h2('在 Karmora 中的应用')
li('爬取 r/NewToReddit 完整 wiki（76 个页面，848KB）')
li('提取新用户友好社区列表（126 个）')
li('提取 Karma 获取指南、红线规则、shadowban 触发条件')

div()

# ── 方法5 ──
h1('方法⑤：Reddit 官方 API（已阻断）')

h2('现状')
b('2023 年起，Reddit 封锁了所有服务器端直接调用 .json 端点的请求（返回 403 Forbidden）。')
b('需要 Client ID + Client Secret 才能使用官方 Data API，但用户没有这些凭证。')

h2('如果未来获得凭证')
li('API 端点：https://oauth.reddit.com/...')
li('Rate limit: 60 请求/分钟（OAuth）')
li('可以获取更完整的数据')

h2('当前状态')
b('在 Karmora 中已搁置。用户提到可以在 developers.mercadolibre.com 注册应用，但 Reddit 的开发者流程不同。')

div()

# ── 方法6 ──
h1('方法⑥：浏览器直接抓取（不稳定）')

h2('原理')
b('直接用 urllib/requests 抓取 Reddit 网页，解析 HTML。')

h2('现状')
b('不可靠。Reddit 使用 Cloudflare 防护，服务器端请求会被 CAPTCHA 拦截。')

h2('例外')
li('Browserbase 等反检测浏览器可以绕过，但成本高')
li('本地 Chrome 通过 CDP（remote-debugging-port=9222）可行，但复杂')

div()

doc.add_page_break()

# ── 扩展方案 ──
h1('社区数据扩展方案')

h2('当前数据规模')
li('社区数量：171 个（40 自采 + 131 官方推荐）')
li('每个社区字段：name, desc, avg_comments, zero_rate, posts_per_day, badge, category')

h2('扩展目标')
li('覆盖 Reddit 上所有主要社区（估计 10,000+）')
li('每个社区增加规则数据（从 rules.json 获取）')
li('每个社区增加详细描述（从 about.json 获取）')

h2('推荐方案：三阶段扩展')

h3('第一阶段：规则库建设（当前可做）')
b('为现有 171 个社区采集规则数据：')
code('// 用户在浏览器 Console 运行')
code('const subs = ["TooAfraidToAsk", "CasualConversation", ...];')
code('const results = {};')
code('for (const sub of subs) {')
code('  const about = await fetch(`/r/${sub}/about.json`, {credentials:"same-origin"});')
code('  const rules = await fetch(`/r/${sub}/rules.json`, {credentials:"same-origin"});')
code('  results[sub] = {')
code('    about: (await about.json()).data,')
code('    rules: (await rules.json()).rules')
code('  };')
code('  await new Promise(r => setTimeout(r, 1000));')
code('}')
code('copy(JSON.stringify(results));  // 复制到剪贴板')

h3('第二阶段：社区发现（需要后端）')
b('通过 PullPush API 或 Reddit 搜索 API 批量发现新社区：')
li('按关键词搜索：/subreddits/search.json?q={keyword}&limit=100')
li('按分类遍历：/subreddits/popular.json, /subreddits/new.json')
li('从已有社区的帖子中提取提到的其他社区')

h3('第三阶段：自动化采集（长期）')
b('建立定时采集管道：')
li('每天自动采集 top 1000 社区的 engagement 数据')
li('每周更新社区规则数据')
li('监控新社区创建')

h2('数据存储建议')
code('communities.json')
code('{')
code('  "TooAfraidToAsk": {')
code('    "name": "TooAfraidToAsk",')
code('    "desc": "匿名提问社区，氛围友善",')
code('    "about": { "subscribers": 500000, "description": "...", ... },')
code('    "rules": [ { "short_name": "...", "description": "..." }, ... ],')
code('    "engagement": { "avg_comments": 11.8, "zero_rate": 2, "posts_per_day": 94 },')
code('    "badge": "excellent",')
code('    "category": "karma_friendly",')
code('    "last_updated": "2026-04-21"')
code('  },')
code('  ...')
code('}')

div()

# ── 快速参考 ──
doc.add_page_break()
h1('快速参考卡片')

h2('采集社区 engagement 数据')
code('// 在 Reddit 页面 Console 运行')
code('const subs = ["sub1", "sub2", ...];')
code('for (const sub of subs) {')
code('  const resp = await fetch(`/r/${sub}/new.json?limit=50`, {credentials:"same-origin"});')
code('  const posts = (await resp.json()).data.children.map(p=>p.data);')
code('  const comments = posts.map(p=>p.num_comments);')
code('  const avg = (comments.reduce((a,b)=>a+b,0)/comments.length).toFixed(1);')
code('  const zero = Math.round(comments.filter(c=>c===0).length/posts.length*100);')
code('  console.log(`${sub}: avg=${avg}, zero=${zero}%`);')
code('  await new Promise(r=>setTimeout(r,500));')
code('}')

h2('采集社区规则')
code('const resp = await fetch("/r/{sub}/rules.json", {credentials:"same-origin"});')
code('const rules = (await resp.json()).rules;')
code('rules.forEach(r => console.log(r.short_name + ": " + r.description));')

h2('采集社区 about')
code('const resp = await fetch("/r/{sub}/about.json", {credentials:"same-origin"});')
code('const about = (await resp.json()).data;')
code('console.log(about.public_description);')
code('console.log("Subscribers:", about.subscribers);')

h2('批量搜索社区')
code('const resp = await fetch("/subreddits/search.json?q={keyword}&limit=100", {credentials:"same-origin"});')
code('const subs = (await resp.json()).data.children.map(s=>s.data.display_name);')
code('console.log(subs);')

h2('从 PullPush 批量采集（Python）')
code('import urllib.request, json, time')
code('for sub in subs:')
code('    url = f"https://api.pullpush.io/reddit/search/submission/?subreddit={sub}&size=50"')
code('    req = urllib.request.Request(url, headers={"User-Agent": "Karmora/1.0"})')
code('    try:')
code('        with urllib.request.urlopen(req, timeout=15) as resp:')
code('            posts = json.loads(resp.read()).get("data", [])')
code('            # 处理数据...')
code('    except: pass')
code('    time.sleep(1)')

out = os.path.expanduser('~/reddit-karma-landing/docs/karmora_data_methods.docx')
doc.save(out)
print(f'Done: {out} ({os.path.getsize(out)/1024:.0f}KB)')
