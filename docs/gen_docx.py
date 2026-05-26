import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(2.5)
    s.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.name = 'PingFang SC'
style.font.size = Pt(11)

def title(t):
    p = doc.add_heading(t, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0x45,0x00); r.font.size = Pt(24)

def sub(t):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(t); r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x78,0x7C,0x7E)

def h1(t):
    p = doc.add_heading(t, level=1)
    for r in p.runs: r.font.color.rgb = RGBColor(0xFF,0x45,0x00); r.font.size = Pt(18)

def h2(t):
    p = doc.add_heading(t, level=2)
    for r in p.runs: r.font.size = Pt(14)

def h3(t):
    p = doc.add_heading(t, level=3)
    for r in p.runs: r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x55,0x55,0x55)

def u(t):
    p = doc.add_paragraph()
    r = p.add_run('用户：'); r.bold = True
    r = p.add_run(f'"{t}"')

def b(t):
    doc.add_paragraph(t).paragraph_format.space_after = Pt(6)

def li(t):
    doc.add_paragraph(t, style='List Bullet')

def div():
    p = doc.add_paragraph()
    r = p.add_run('─'*60); r.font.color.rgb = RGBColor(0xDD,0xDD,0xDD); r.font.size = Pt(8)

title('Karmora 项目完整开发日志')
sub('从想法诞生到产品成型的完整时间线记录')
sub('生成日期：2026-04-21')
doc.add_page_break()

h1('2026-04-18 — 产品想法诞生')
h2('对话内容')
u('我在运营reddit时，遇到了karma低，被社区下架帖子的情况，并且在找了一圈小红书上对reddit的运营的难处，都在说很难涨karma，我现在觉得这是一个变现逻辑')
h2('产生了什么')
li('产品想法：Reddit Karma 增长是一个真实痛点')
li('方向 A：Reddit Karma 增长 SaaS 工具')
li('市场调研：竞品分析（Postpone $100/月、Later、Buffer 等）')
li('关键发现：所有现有工具都假设你"已经有能发帖的账号"，没人解决"新号冷启动"问题')
div()

h2('Landing Page 设计')
u('详细了解方向A')
u('Landing page里面需要什么功能')
u('直接帮我做一个landing page')
b('产生了：~/reddit-karma-landing/index.html')
div()

h2('技术调研')
u('我获取不到reddit的开发者client ID和secret')
u('我在reddit的开发者中心可以创建那个好像叫devort的应用，这个是可以干嘛的')
li('关键约束：没有 Reddit API 凭证')
li('调研发现：Chrome 扩展可以借用用户已登录身份进行同源请求')
li('关键技术决策：Chrome 扩展 + PullPush API 方案')
div()

h2('小红书推广')
u('我现在准备把landing page发到小红书上，看看是什么反应，给我生成一个小红书的文案')
li('3 个版本的小红书文案（痛点共鸣型、功能展示型、对比型）')
div()

h2('Chrome 扩展开发 v0.1.0')
b('创建了完整的 Chrome 扩展，共 10 个文件：')
li('manifest.json / content.js(910行) / content.css / popup.js / popup.html / background.js / subreddit-db.js / cache-manager.js / content-assistant.js / path-planner.js')
b('功能清单：')
li('1. 用户 karma + 账号年龄显示')
li('2. 社区实时分析面板')
li('3. 社区数据库（40个社区）')
li('4. 安全监控（发帖拦截）')
li('5. AI 内容助手（模板字符串替换）')
li('6. 路径规划')
div()

h2('账号年龄问题')
u('我的账号karma是1，插件上显示的不对啊')
u('账号年龄显示的是"？"，你要去解决这个')
u('我觉得这个功能，可以丢掉，改为用户手动填写')
u('注册时间的查询方法：点击右上角的头像→view profile→ACHIEVEMENTS→View All→Getting Started里的徽章可以看到日期')
li('解决：自动获取不可靠 → 改为手动填写 + 查询指引')
div()

h2('插件/Web端架构决策')
u('我认为有一些功能应该放在网页上')
u('插件上放太多功能可能会太臃肿')
li('架构决策：Chrome 扩展精简，Web Dashboard 承载复杂功能')
li('Web Dashboard 5 页面：首页/社区数据库/智能匹配/路径规划/AI 助手')
div()

h2('想法1诞生')
u('在landing page里加入一个问答功能，让用户说出他们的需求，然后后台搜索数据库（reddit上所有的社区的数据），语义匹配推荐合适的subreddit')
u('这个想法帮我命名为想法1（待实现）')
li('想法1：用户输入中文需求 → 后台搜索 Reddit 全部社区 → 语义匹配推荐 subreddit')
li('需要：后端 + 向量搜索/语义匹配')
div()

doc.add_page_break()
h1('2026-04-19 — 功能调试日')
u('截止到目前这个项目的功能我需要都调试一遍，请你列出每一个功能，给我调试指南')
b('列出了 10 个功能点，逐个调试。')
div()

h3('功能1：Karma + 账号显示')
u('当我离开了我的账号页面，插件的页面就会脱离，你认为这个可行吗')
b('解决：新增从 Reddit header DOM 抓取 karma，不依赖个人页面')
div()

h3('功能3：社区数据库')
u('打开了某一个subreddit，一直显示正在分析')
u('当我打开某一个subreddit的时候，在插件上应该显示这个subreddit的中文简介')
b('解决：添加中文简介显示，移除内容助手')
div()

h3('功能4/6：安全监控')
u('安全检查的条件是什么')
u('安全提醒窗口确实跳出来了，但是下面的返回和仍然发帖按钮按不动')
u('文字没有完全显示在按钮的框内')
b('解决：多次修复弹窗 → 用 dialog 替代 div z-index')
b('触发规则：URL变化检测（/submit→其他=提交成功）')
div()

h3('Karma 数乱跳')
u('当我点开某个subreddit的时候，总karma数会乱跳')
u('我不管你怎么修这个代码，我需要的最终结果是，插件上的总karma数要和个人中心的一致')
b('解决：SPA导航导致重复获取 → 全局缓存 _cachedUserData')
div()

doc.add_page_break()
h1('2026-04-20 — Web Dashboard 开发日')
u('今天开始功能8')
u('等会，你先把项目梳理一遍')
div()

h2('数据采集')
u('我需要知道这个新手推荐的逻辑是什么，有什么数据支撑这些社区适合新手吗')
u('你能获取到reddit上的所有社区吗，只是获取社区名字这么简单')
li('使用 PullPush API 采集 40 个社区的真实数据')
li('新增字段：avg_comments / zero_rate / posts_per_day')
div()

h2('Landing Page 改版')
u('现在应该对landing page做一个大改动，需要有首页和其他功能按钮')
b('从单页面 → 5 个独立 HTML 页面 + 共享导航栏')
div()

h2('F1-F15 问题清单')
u('你先把新加的功能（半成品和缺失）全部列出来，给这些功能编号')
li('F1: 首页 CTA / F2: 数据数量不匹配 / F3: 缺少排序 / F4: 缺少搜索 / F5: 卡片无链接')
li('F6: 搜索范围太窄(暂缓) / F7: 数据不同步 / F8: 输入太简单 / F9: 计划太通用 / F10: 只有9个社区')
li('F11: 按钮无反馈 / F12: 只是字符串替换 / F13: 复制按钮 / F14: 发帖警告 / F15: CTA')
div()

h2('F1修复（首页CTA）')
u('把流动的效果去掉吧，就留下上面的大字和下面的按钮吧')
b('首页 CTA 简化：大标题+按钮')
div()

h2('F2修复（数据数量）')
b('补充 5 个缺失社区（35→40）')
div()

h2('F3修复（排序）')
u('你定义的那三个排序是以什么依据为排序的')
u('可以把这三种排序方法都加上去，然后让用户自己选择')
b('zero_rate / avg_comments / posts_per_day，支持升降序切换')
div()

h2('F4修复（搜索）')
u('不要去掉一开始的那个版本，在那个版本的基础上加上搜索功能')
b('搜索=筛选当前视图，不替换')
div()

h2('F7-F10修复')
li('F7: 创建 js/data.js 作为共享数据源（171 个社区）')
li('F8: 路径规划添加 karma+年龄输入')
li('F9: 基于 r/NewToReddit wiki 研究重新设计（年龄逻辑+红线规则）')
li('F10: AI 助手下拉菜单 9→171 个社区')
div()

doc.add_page_break()
h1('2026-04-21 — Wiki 分析 + 产品审视日')
h2('Wiki 调研')
u('需要先弄清楚，路径是否真的可以增长karma，需要reddit官方指南')
u('你写一个爬虫给我，我需要把r/NewToReddit的wiki的所有信息都爬出来')
u('获取官方推荐的新人友好社区，根据调研结果重新制定路径规划')
b('Wiki 数据：r/NewToReddit 完整 wiki（76 个页面，848KB）')
div()

h2('Wiki 深度分析')
u('我需要你再去重新分析一遍之前的那个wiki')
b('关键发现：')
li('AI 内容会被踩："many people consider AI generated text to be low effort"')
li('评论比发帖更容易获得 karma')
li('shadowban 机制（太快/重复/外链/多账号）')
li('126 个官方认证新用户友好社区')
li('CQS（贡献者质量评分）系统')
div()

h2('产品审视')
u('那看完你的严重程度之后，我是不是要对项目大改')
u('OK，那其实后期也是只需要对这个项目进行小修小补，我理解的意思对吗')
b('核心矛盾：产品生成 AI 内容，但 Reddit 用户会踩 AI 内容')
b('定位转变：从"生成"转向"校验"')
b('结论：不需要大改，小修小补')
div()

h2('LLM Agent 讨论')
u('这四个核心需求首先需要选择一款LLM...以接入的LLM为核心')
u('要让LLM遵守规则，相当于是让以LLM为基础建立一个agent')
u('这一步先放着')
li('LLM 选型：DeepSeek V3 / GPT-4o-mini / Claude / Gemini / Qwen')
li('Agent 架构：LLM + 规则库 + 目标指令 + 校验护栏')
li('搁置：LLM Agent 和 F15')
div()

h2('F11修复')
b('内容类型按钮 → .active class：橙色背景+白字+加粗+阴影')
div()

h2('F12修复')
li('中译英：MyMemory API（免费）')
li('模块化模板：960 种组合/类型 × 4 类型 = 3,840 种输出')
li('新功能：翻译结果显示、重新生成按钮、社区警告')
div()

h2('数据整理')
b('备份：~/.hermes/backups/karmora_20260421_164804/（72 个文件，33.9MB）')
li('Memory: 更新 Karmora 条目，保留想法1')
li('Sessions: 64→22 个（释放 20.8MB）')
li('Skills: 3→2 个（合并）')
li('Cache: 3→1 个（删除过时文件）')

out = os.path.expanduser('~/reddit-karma-landing/docs/karmora_full_timeline.docx')
doc.save(out)
print(f'Done: {out} ({os.path.getsize(out)/1024:.0f}KB)')
